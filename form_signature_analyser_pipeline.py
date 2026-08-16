import os
import json
import csv
from typing import List, Tuple
from PIL import Image, ImageDraw
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from schemas import Stage1SignatureAnalysis, Stage2DatasetComparison

class FormSignatureAnalyserPipeline:
    """
    End-to-End Application Form Signature Analysis & Reference Dataset Comparator
    using Google Gemini (Multimodal Vision API).
    """

    def __init__(self, api_key: str = None, model_name: str = "gemini-2.5-flash"):
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        if not self.api_key:
            raise ValueError("GEMINI_API_KEY environment variable required.")
        self.client = genai.Client(api_key=self.api_key)
        self.model_name = model_name

    def analyze_form(self, form_image_path: str) -> Tuple[Stage1SignatureAnalysis, Image.Image, Image.Image]:
        """
        Stage 1: Analyzes application form for signature presence, classification, bounding box, and metadata.
        Crops signature and draws bounding box overlay.
        """
        img = Image.open(form_image_path).convert("RGB")

        prompt = """
        Analyze this application form image for handwritten signatures, initials, or digital stamps.
        Determine:
        1. Whether a signature is present.
        2. Classification: 'wet_signature' (handwritten on paper), 'pasted_image' (digitally inserted), 'signature_does_not_exist', or 'uncertain'.
        3. Confidence score (0.0 to 1.0).
        4. Extracted signatory name, role, and date if visible.
        5. Bounding box coordinates in normalized scale [ymin, xmin, ymax, xmax] from 0 to 1000.
        6. Explanation, visual evidence bullet points, and image limitations.
        Adhere strictly to the JSON schema.
        """

        response = self.client.models.generate_content(
            model=self.model_name,
            contents=[img, prompt],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=Stage1SignatureAnalysis,
                temperature=0.1,
            ),
        )

        analysis = Stage1SignatureAnalysis.model_validate_json(response.text)

        # Draw bounding box & crop signature if present
        img_width, img_height = img.size
        annotated_img = img.copy()
        draw = ImageDraw.Draw(annotated_img)

        bbox = analysis.bounding_box
        ymin = int(bbox.ymin * img_height / 1000.0)
        xmin = int(bbox.xmin * img_width / 1000.0)
        ymax = int(bbox.ymax * img_height / 1000.0)
        xmax = int(bbox.xmax * img_width / 1000.0)

        # Ensure box is valid
        xmin, ymin = max(0, xmin), max(0, ymin)
        xmax, ymax = min(img_width, max(xmin + 10, xmax)), min(img_height, max(ymin + 10, ymax))

        draw.rectangle([xmin, ymin, xmax, ymax], outline="red", width=4)
        cropped_sig = img.crop((xmin, ymin, xmax, ymax))

        return analysis, annotated_img, cropped_sig

    def compare_with_reference(self, cropped_sig_img: Image.Image, ref_sig_path: str) -> Stage2DatasetComparison:
        """
        Stage 2: Compares cropped signature against a single reference dataset signature image.
        """
        ref_img = Image.open(ref_sig_path).convert("RGB")
        filename = os.path.basename(ref_sig_path)

        prompt = f"""
        Compare Image 1 (Cropped signature from form) against Image 2 (Reference dataset signature '{filename}').
        Analyze:
        - Structural trajectory, initial strokes, character loops, terminal strokes, and flourishes.
        - Determine match_result: 'PERFECT MATCH', 'PARTIAL MATCH', 'NO MATCH', or 'UNCERTAIN'.
        - Provide confidence score (integer 1-100).
        - Provide detailed comparative explanation, bulleted visual evidence, and limitations.
        Do NOT rely on filenames; base your comparison purely on visual characteristics.
        Adhere strictly to the JSON schema.
        """

        response = self.client.models.generate_content(
            model=self.model_name,
            contents=[cropped_sig_img, ref_img, prompt],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=Stage2DatasetComparison,
                temperature=0.1,
            ),
        )

        result = Stage2DatasetComparison.model_validate_json(response.text)
        result.dataset_signature_name = filename
        return result

    def generate_docx_report(
        self,
        output_docx_path: str,
        stage1_analysis: Stage1SignatureAnalysis,
        annotated_form_path: str,
        cropped_sig_path: str,
        comparison_results: List[Tuple[Stage2DatasetComparison, str]]
    ):
        """
        Generates a landscape Word (.docx) report formatted with tables and comparison side-by-sides.
        """
        doc = Document()
        
        # Set landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        # Title
        title = doc.add_heading("Form Signature Analysis & Reference Comparison Report", level=1)
        
        # Section 1: Stage 1 Analysis
        doc.add_heading("1. Application Form Signature Analysis", level=2)
        
        p = doc.add_paragraph()
        p.add_run(f"Signature Present: ").bold = True
        p.add_run(f"{stage1_analysis.signature_present}\n")
        p.add_run(f"Classification: ").bold = True
        p.add_run(f"{stage1_analysis.classification}\n")
        p.add_run(f"Detection Confidence: ").bold = True
        p.add_run(f"{stage1_analysis.confidence * 100:.1f}%\n")
        p.add_run(f"Signatory Name: ").bold = True
        p.add_run(f"{stage1_analysis.signatory_name or 'N/A'}\n")
        p.add_run(f"Explanation: ").bold = True
        p.add_run(f"{stage1_analysis.explanation}")

        # Add images
        doc.add_paragraph("Form Bounding Box & Signature Crop:")
        tbl_img = doc.add_table(rows=1, cols=2)
        cell_1 = tbl_img.cell(0, 0)
        cell_2 = tbl_img.cell(0, 1)
        cell_1.paragraphs[0].add_run().add_picture(annotated_form_path, width=Inches(3.5))
        cell_2.paragraphs[0].add_run().add_picture(cropped_sig_path, width=Inches(2.5))

        # Section 2: Stage 2 Dataset Comparison
        doc.add_heading("2. Reference Signature Dataset Comparison", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Cropped Signature"
        hdr_cells[1].text = "Dataset Signature"
        hdr_cells[2].text = "Comparison Result & Evidence"

        for comp, ref_img_path in comparison_results:
            row_cells = table.add_row().cells
            
            # Cell 0: Cropped signature
            p0 = row_cells[0].paragraphs[0]
            p0.add_run().add_picture(cropped_sig_path, width=Inches(1.8))
            
            # Cell 1: Reference dataset image + info
            p1 = row_cells[1].paragraphs[0]
            p1.add_run().add_picture(ref_img_path, width=Inches(1.8))
            p1.add_run(f"\nFile: {comp.dataset_signature_name}\nPath: {os.path.abspath(ref_img_path)}")

            # Cell 2: Match result & analysis
            p2 = row_cells[2].paragraphs[0]
            r_match = p2.add_run(f"Result: {comp.match_result} ({comp.confidence}/100)\n")
            r_match.bold = True
            p2.add_run(f"Explanation: {comp.explanation}\n\nEvidence:\n")
            for ev in comp.visual_evidence:
                p2.add_run(f"• {ev}\n")

        doc.save(output_docx_path)
        print(f"📄 DOCX report generated: {output_docx_path}")
